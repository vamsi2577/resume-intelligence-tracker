"""
DOCX generation utility for resume generation.

Converts a ResumeRequest into a .docx BytesIO stream.
Supports **bold** markdown syntax in bullets and summary text.
Zero FastAPI / DB dependencies — pure document generation.
"""
from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from src.schemas.resume_generator import ResumeRequest
from src.utils.logger import get_logger

logger = get_logger(__name__)


# ── Helpers ───────────────────────────────────────────────

def sanitize_filename(text: str) -> str:
    text = text.strip()
    text = re.sub(r'[^\w\s-]', '', text)
    text = re.sub(r'[-\s]+', '_', text)
    return text


def _add_bottom_border(paragraph) -> None:
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _format_header(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    _add_bottom_border(p)


def _add_markdown_paragraph(doc: Document, text: str, style: str | None = None):
    """
    Parses **bold** markdown and adds styled runs to a paragraph.
    """
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = p.add_run(part)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        if i % 2 == 1:
            run.bold = True
    return p


# ── Main builder ──────────────────────────────────────────

def build_docx(data: ResumeRequest) -> io.BytesIO:
    """
    Generates a .docx BytesIO stream from a ResumeRequest.
    Raises ValueError if required sections are missing.
    """
    doc = Document()

    # Margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # ── Name header ───────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(data.full_name.upper())
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(data.contact_info)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'

    # ── Summary ───────────────────────────────────────────
    if data.summary:
        _format_header(doc, "Summary")
        p = _add_markdown_paragraph(doc, data.summary.summary_text)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for point in data.summary.summary_points:
            _add_markdown_paragraph(doc, point, style='List Bullet')

    # ── Skills ────────────────────────────────────────────
    if data.skills:
        _format_header(doc, "Technical Skills")
        table = doc.add_table(rows=0, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(2.2)
        table.columns[1].width = Inches(5.3)

        for skill_cat in data.skills:
            row_cells = table.add_row().cells
            p = row_cells[0].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(skill_cat.category)
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

            p = row_cells[1].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(", ".join(skill_cat.items))
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

    # ── Experience ────────────────────────────────────────
    if data.experience:
        _format_header(doc, "Experience")
        for job in data.experience:
            p = doc.add_paragraph()
            p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)

            run = p.add_run(f"{job.company} – {job.title}")
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(12)

            run = p.add_run(f"\t{job.date}")
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(12)

            for bullet in job.bullets:
                p = _add_markdown_paragraph(doc, bullet, style='List Bullet')
                p.paragraph_format.space_after = Pt(1)

            if job.tools:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                run = p.add_run("Tools & Technologies: ")
                run.bold = True
                run.italic = True
                run.font.size = Pt(11)
                run.font.name = 'Calibri'
                run = p.add_run(", ".join(job.tools))
                run.italic = True
                run.font.size = Pt(11)
                run.font.name = 'Calibri'

    # ── Certifications ────────────────────────────────────
    if data.certifications:
        _format_header(doc, "Certifications")
        for cert in data.certifications:
            p = doc.add_paragraph(cert.name, style='List Bullet')
            p.runs[0].font.name = 'Calibri'
            p.runs[0].font.size = Pt(11)

    # ── Education ─────────────────────────────────────────
    if data.education:
        _format_header(doc, "Education")
        for edu in data.education:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(f"{edu.degree} – {edu.university}")
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

            # add_paragraph(text) emits no runs when `text` is empty (Pydantic
            # allows "" for edu.details, and the LLM sometimes returns blank
            # for compressed entries). Build the run explicitly so styling
            # always has a target.
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            if edu.details:
                run = p.add_run(edu.details)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    logger.info(
        "DOCX generated",
        extra={"company": data.target_company, "job_title": data.job_title},
    )
    return stream


def build_filename(data: ResumeRequest) -> str:
    """Returns a sanitized filename for the generated DOCX."""
    parts = data.full_name.strip().split()
    lastname = sanitize_filename(parts[-1]) if parts else "Candidate"
    company = sanitize_filename(data.target_company)
    return f"{company}_{lastname}_Resume.docx"
